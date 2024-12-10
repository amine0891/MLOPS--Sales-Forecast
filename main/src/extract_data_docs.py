import re
import pandas as pd
import os
import docx

# nettoyage des données des caractères spéciaux
unwanted_chars = ['\n', '\t', '\xa0', 'â\x80\x93']


def clean_up_text(text, unwanted_chars=unwanted_chars):
    for char in unwanted_chars:
        text = text.replace(char, '').lower()

    return text


# nettoyage des chemins
def clean_up_path(path):
    path = re.sub(r'\b(RP|1 Trans|2 Trans|EXPORT)\b', '', path)
    path = path.strip().split()[:2]
    path[1] = re.sub(r'.docx', '', path[1])
    path[1] = re.sub(r'[^A-Za-zéèôû]', '', path[1])
    return ' '.join(path)


# Créer la liste des pays et codes de produits à partir de deux fichiers CodeProduitDataBase et Pays DataBase

def create_lists():
    Pays = []
    CodeProduit = []
    with open('../../data/Countries.txt', 'r', encoding="utf-8") as file1:
        for line in file1:
            Pays.append(clean_up_text(line))
    Pays = set(Pays)
    with open('../../data/Product-codes.txt', 'r', encoding="utf-8") as file2:
        for line in file2:
            CodeProduit.append(clean_up_text(line))
    return Pays, CodeProduit


def append_to_df(root, file, data, record):
    tmp = clean_up_path(file)
    date = tmp + " " + root.split("\\")[1]
    for p in record[0]:
        for i in range(p[1]):
            data.append([p[0], record[1], date])


def create_excel_sales_data():
    directory = "../../data/Sales data"
    data = []
    Pays, CodeProduit = create_lists()
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.startswith('~'):
                continue
            if file.endswith('.docx'):
                path = os.path.join(root, file).replace("\\", "/")
                try:
                    doc = docx.Document(path)
                except Exception:
                    continue
                record = [None, None, None]
                saved_ids = []
                for i in range(len(doc.paragraphs)):
                    paragraph = clean_up_text(doc.paragraphs[i].text)
                    pp = re.sub(r"[ \(\)]", "", paragraph)
                    for id in CodeProduit:
                        tid = id.replace(" ", '')
                        match = re.search(tid, pp)
                        if match:
                            if record[0] is not None:
                                append_to_df(root, file, data, record)
                                record = [None, None, None]
                                saved_ids = []
                            position = match.start()
                            n = 1
                            if position > 1 and pp[position - 1] == '*':
                                if pp[position - 3].isdigit():
                                    n = int(pp[position - 3] + pp[position - 2])
                                else:
                                    n = int(pp[position - 2])
                            saved_ids.append((id, n))
                            pp = re.sub(rf'{re.escape(tid)}', '', pp)
                    record[0] = saved_ids

                    for pay in Pays:
                        if paragraph.find(pay.lower()) != -1:
                            record[1] = pay
                    if i == len(doc.paragraphs) - 1 and record[0] is not None:
                        append_to_df(root, file, data, record)

    df = pd.DataFrame(data, columns=['CodeProduit', 'Pays', 'Date'])
    df.to_excel("../../data/Sales_data.xlsx", index=False)
    print('data successfully extracted from docs !')


if __name__=='__main__':
    create_excel_sales_data()
